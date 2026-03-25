from dataclasses import dataclass, field
from pathlib import Path
from docx import Document
from pipelines.ingestion.extractors.pdf_extractor import ExtractedPage


async def extract_docx(file_path: str) -> list[ExtractedPage]:
    pages: list[ExtractedPage] = []
    doc = Document(file_path)
    path = Path(file_path)

    current_section = ""
    current_level = None
    buffer = []
    page_num = 1

    def flush_buffer():
        nonlocal buffer, page_num
        if buffer:
            pages.append(ExtractedPage(
                page_number=page_num,
                content="\n".join(buffer),
                content_type="text",
                metadata={
                    "section": current_section,
                    "heading_level": current_level,
                    "source_file": path.name,
                },
            ))
            page_num += 1
            buffer = []

    for para in doc.paragraphs:
        style_name = para.style.name.lower()

        if style_name.startswith("heading"):
            flush_buffer()
            level = style_name.replace("heading ", "").strip()
            current_section = para.text
            current_level = f"h{level}"
            buffer.append(f"{'#' * int(level)} {para.text}")
        else:
            if para.text.strip():
                buffer.append(para.text)

    flush_buffer()

    # Extract tables
    for table_idx, table in enumerate(doc.tables):
        rows = []
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
            if row_idx == 0:
                rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
        pages.append(ExtractedPage(
            page_number=page_num,
            content="\n".join(rows),
            content_type="table",
            metadata={"table_index": table_idx, "source_file": path.name},
        ))
        page_num += 1

    return pages
